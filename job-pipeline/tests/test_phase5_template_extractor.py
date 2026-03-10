"""Phase 5 tests: Template extractor."""

import json
import shutil
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt
from lxml import etree

from agent.template_extractor import (
    unpack_docx,
    find_bullet_nodes,
    save_template_map,
    load_template_map,
    get_paragraph_text,
    is_bullet_paragraph,
    is_section_header,
)


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    dirpath = tempfile.mkdtemp()
    yield Path(dirpath)
    shutil.rmtree(dirpath)


@pytest.fixture
def minimal_docx(temp_dir):
    """Create a minimal DOCX with one employer (5 bullets) and one project (3 bullets)."""
    doc = Document()
    
    # Work Experience section
    doc.add_heading("Work Experience", level=1)
    
    # Company header
    doc.add_paragraph("Jaguar TCS Racing")
    doc.add_paragraph("Software Engineer | London | 2022-Present")
    
    # Bullets with ▪ character
    bullets = [
        "▪ Designed REST API for streaming live race telemetry from WinTax",
        "▪ Developed full-stack ReactJS data visualisation apps for time-critical race data",
        "▪ Collaborate in ML/AI pipeline planning alongside data scientists",
        "▪ Deployed CI/CD pipeline for trackside executables ensuring software reliability",
        "▪ Create energy regeneration analysis tool allowing strategists to make informed calls",
    ]
    for bullet in bullets:
        doc.add_paragraph(bullet)
    
    # Technical Projects section
    doc.add_heading("Technical Projects", level=1)
    
    # Project header
    doc.add_paragraph("Formula Student Lap Time Simulator")
    
    # Project bullets
    project_bullets = [
        "▪ Designed modular steady-state lap time sim using point-mass and bicycle models",
        "▪ Combined physics-based and empirical models to distill complex racing dynamics",
        "▪ Created tyre model with Pacejka magic formula ensuring accuracy",
    ]
    for bullet in project_bullets:
        doc.add_paragraph(bullet)
    
    # Save
    docx_path = temp_dir / "test_cv.docx"
    doc.save(str(docx_path))
    return docx_path


@pytest.fixture
def multi_subsection_docx(temp_dir):
    """Create a DOCX with multiple employers and projects."""
    doc = Document()
    
    # Work Experience section
    doc.add_heading("Work Experience", level=1)
    
    # First company
    doc.add_paragraph("Jaguar TCS Racing")
    doc.add_paragraph("Software Engineer | London")
    doc.add_paragraph("▪ Designed REST API for streaming live race telemetry")
    doc.add_paragraph("▪ Developed full-stack ReactJS data visualisation apps")
    doc.add_paragraph("▪ Deployed CI/CD pipeline for trackside executables")
    
    # Second company
    doc.add_paragraph("Republic of Singapore Navy")
    doc.add_paragraph("Technician | Singapore")
    doc.add_paragraph("▪ Trained on Singapore Navy's largest marine vessel")
    doc.add_paragraph("▪ Led remote war-time firefighting simulation")
    doc.add_paragraph("▪ Contributed to 50 pre- and post-mission debriefs")
    
    # Technical Projects section  
    doc.add_heading("Technical Projects", level=1)
    
    # First project
    doc.add_paragraph("Formula Student Lap Time Simulator")
    doc.add_paragraph("▪ Designed modular steady-state lap time sim")
    doc.add_paragraph("▪ Combined physics-based and empirical models")
    doc.add_paragraph("▪ Created tyre model with Pacejka magic formula")
    
    # Second project
    doc.add_paragraph("2D CFD Radiator Optimisation Simulator")
    doc.add_paragraph("▪ Developed a 2D incompressible fluid solver")
    doc.add_paragraph("▪ Automated batch simulations and visualisations")
    doc.add_paragraph("▪ Leveraged parallel processing for optimisation")
    
    # Save
    docx_path = temp_dir / "multi_cv.docx"
    doc.save(str(docx_path))
    return docx_path


class TestUnpackDocx:
    """Tests for unpack_docx function."""
    
    def test_produces_document_xml(self, minimal_docx, temp_dir):
        """unpack_docx produces directory with word/document.xml."""
        output_dir = temp_dir / "unpacked"
        doc_xml_path = unpack_docx(minimal_docx, output_dir)
        
        assert doc_xml_path.exists()
        assert doc_xml_path.name == "document.xml"
        assert doc_xml_path.parent.name == "word"
    
    def test_creates_output_directory(self, minimal_docx, temp_dir):
        """unpack_docx creates the output directory."""
        output_dir = temp_dir / "new_output"
        assert not output_dir.exists()
        
        unpack_docx(minimal_docx, output_dir)
        
        assert output_dir.exists()
        assert (output_dir / "word").exists()
    
    def test_raises_on_missing_docx(self, temp_dir):
        """unpack_docx raises FileNotFoundError on missing file."""
        fake_path = temp_dir / "nonexistent.docx"
        output_dir = temp_dir / "output"
        
        with pytest.raises(FileNotFoundError):
            unpack_docx(fake_path, output_dir)
    
    def test_overwrites_existing_output(self, minimal_docx, temp_dir):
        """unpack_docx cleans existing output directory."""
        output_dir = temp_dir / "unpacked"
        output_dir.mkdir()
        (output_dir / "old_file.txt").write_text("old content")
        
        unpack_docx(minimal_docx, output_dir)
        
        assert not (output_dir / "old_file.txt").exists()
        assert (output_dir / "word" / "document.xml").exists()


class TestFindBulletNodes:
    """Tests for find_bullet_nodes function."""
    
    def test_returns_dict_with_both_sections(self, minimal_docx, temp_dir):
        """find_bullet_nodes returns dict with both sections."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        
        result = find_bullet_nodes(doc_xml)
        
        assert 'work_experience' in result
        assert 'technical_projects' in result
    
    def test_finds_subsections(self, multi_subsection_docx, temp_dir):
        """At least 2 subsections in work_experience, at least 1 in technical_projects."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(multi_subsection_docx, output_dir)
        
        result = find_bullet_nodes(doc_xml)
        
        assert len(result['work_experience']) >= 2, "Should have at least 2 employers"
        assert len(result['technical_projects']) >= 1, "Should have at least 1 project"
    
    def test_finds_bullets_per_subsection(self, minimal_docx, temp_dir):
        """At least 3 bullets found per subsection."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        
        result = find_bullet_nodes(doc_xml)
        
        for section, subsections in result.items():
            for name, data in subsections.items():
                bullets = data.get('bullet_xpaths', [])
                assert len(bullets) >= 3, f"{section}/{name} should have at least 3 bullets, got {len(bullets)}"
    
    def test_bullet_xpaths_are_strings(self, minimal_docx, temp_dir):
        """Bullet XPaths should be valid strings."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        
        result = find_bullet_nodes(doc_xml)
        
        for section, subsections in result.items():
            for name, data in subsections.items():
                for xpath in data.get('bullet_xpaths', []):
                    assert isinstance(xpath, str)
                    assert xpath.startswith('/'), f"XPath should be absolute: {xpath}"


class TestOriginalDocxUnmodified:
    """Tests ensuring original DOCX is never modified."""
    
    def test_original_docx_unmodified(self, minimal_docx, temp_dir):
        """Original DOCX unmodified after extraction."""
        # Get original hash
        original_content = minimal_docx.read_bytes()
        original_size = len(original_content)
        
        # Run extraction
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        find_bullet_nodes(doc_xml)
        
        # Verify unchanged
        new_content = minimal_docx.read_bytes()
        assert len(new_content) == original_size
        assert new_content == original_content


class TestTemplateMapJson:
    """Tests for template_map.json output."""
    
    def test_valid_json_output(self, minimal_docx, temp_dir):
        """template_map.json is valid JSON."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        result = find_bullet_nodes(doc_xml)
        
        map_path = temp_dir / "template_map.json"
        save_template_map(result, map_path)
        
        # Should parse without error
        with open(map_path, 'r', encoding='utf-8') as f:
            loaded = json.load(f)
        
        assert 'work_experience' in loaded
        assert 'technical_projects' in loaded
    
    def test_load_template_map(self, minimal_docx, temp_dir):
        """load_template_map correctly loads saved map."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        result = find_bullet_nodes(doc_xml)
        
        map_path = temp_dir / "template_map.json"
        save_template_map(result, map_path)
        
        loaded = load_template_map(map_path)
        assert loaded == result
    
    def test_idempotent_extraction(self, minimal_docx, temp_dir):
        """Running twice produces identical output."""
        output_dir = temp_dir / "unpacked"
        map_path = temp_dir / "template_map.json"
        
        # First run
        doc_xml = unpack_docx(minimal_docx, output_dir)
        result1 = find_bullet_nodes(doc_xml)
        save_template_map(result1, map_path)
        content1 = map_path.read_text()
        
        # Second run
        doc_xml = unpack_docx(minimal_docx, output_dir)
        result2 = find_bullet_nodes(doc_xml)
        save_template_map(result2, map_path)
        content2 = map_path.read_text()
        
        assert content1 == content2


class TestHelperFunctions:
    """Tests for helper functions."""
    
    def test_is_bullet_paragraph_with_small_square(self):
        """Detects ▪ (U+25AA) bullet."""
        assert is_bullet_paragraph("▪ This is a bullet")
        assert is_bullet_paragraph("  ▪ Indented bullet")
    
    def test_is_bullet_paragraph_with_black_square(self):
        """Detects ▪ (U+25A0) bullet."""
        assert is_bullet_paragraph("▪ This is a bullet")
    
    def test_is_bullet_paragraph_with_dash(self):
        """Detects - bullet."""
        assert is_bullet_paragraph("- This is a bullet")
    
    def test_is_bullet_paragraph_rejects_non_bullet(self):
        """Rejects non-bullet text."""
        assert not is_bullet_paragraph("Regular paragraph")
        assert not is_bullet_paragraph("Jaguar TCS Racing")
        assert not is_bullet_paragraph("")

        def test_is_bullet_paragraph_detects_word_numpr(self):
                """Detects Word list bullets via w:numPr even without visible bullet character."""
                para_xml = """
                <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:pPr>
                        <w:numPr>
                            <w:ilvl w:val="0"/>
                            <w:numId w:val="1"/>
                        </w:numPr>
                    </w:pPr>
                    <w:r><w:t>Designed API layer for telemetry pipeline</w:t></w:r>
                </w:p>
                """
                para_element = etree.fromstring(para_xml)
                assert is_bullet_paragraph("Designed API layer for telemetry pipeline", para_element)
    
    def test_is_section_header_work_experience(self):
        """Detects Work Experience header."""
        assert is_section_header("Work Experience") == "work_experience"
        assert is_section_header("WORK EXPERIENCE") == "work_experience"
        assert is_section_header("Professional Experience") == "work_experience"
    
    def test_is_section_header_technical_projects(self):
        """Detects Technical Projects header."""
        assert is_section_header("Technical Projects") == "technical_projects"
        assert is_section_header("Projects") == "technical_projects"
        assert is_section_header("Personal Projects") == "technical_projects"
    
    def test_is_section_header_non_header(self):
        """Non-headers return None."""
        assert is_section_header("Jaguar TCS Racing") is None
        assert is_section_header("Software Engineer") is None
        assert is_section_header("▪ A bullet point") is None


class TestBulletCounts:
    """Tests for expected bullet counts."""
    
    def test_minimal_docx_bullet_count(self, minimal_docx, temp_dir):
        """Minimal DOCX should have 5 work bullets and 3 project bullets."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(minimal_docx, output_dir)
        result = find_bullet_nodes(doc_xml)
        
        work_bullets = sum(len(s['bullet_xpaths']) for s in result['work_experience'].values())
        project_bullets = sum(len(s['bullet_xpaths']) for s in result['technical_projects'].values())
        
        assert work_bullets == 5, f"Expected 5 work bullets, got {work_bullets}"
        assert project_bullets == 3, f"Expected 3 project bullets, got {project_bullets}"
    
    def test_multi_subsection_bullets(self, multi_subsection_docx, temp_dir):
        """Multi-subsection DOCX should have correct bullet counts."""
        output_dir = temp_dir / "unpacked"
        doc_xml = unpack_docx(multi_subsection_docx, output_dir)
        result = find_bullet_nodes(doc_xml)
        
        total_work = sum(len(s['bullet_xpaths']) for s in result['work_experience'].values())
        total_projects = sum(len(s['bullet_xpaths']) for s in result['technical_projects'].values())
        
        assert total_work == 6, f"Expected 6 work bullets, got {total_work}"
        assert total_projects == 6, f"Expected 6 project bullets, got {total_projects}"
