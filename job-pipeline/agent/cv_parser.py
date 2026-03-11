"""Extract raw text sections from uploaded CV files.

Supports .docx (via python-docx) and .pdf (via pdfplumber).
Returns a deterministic list of RawSection dicts ready for the UI
section-reviewer. No LLM is used here — headings are detected purely
by heuristics so the detection step is instant and token-free.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

# -----------------------  known section labels  -----------------------

SectionType = Literal[
    "work_experience",
    "technical_projects",
    "education",
    "skills",
    "summary",
    "cover_letter",
    "other",
]

# Maps normalised heading text → detected section type
_HEADING_PATTERNS: list[tuple[str, SectionType]] = [
    # Work experience variants
    (r"work\s+experience", "work_experience"),
    (r"professional\s+experience", "work_experience"),
    (r"employment(\s+history)?", "work_experience"),
    (r"experience", "work_experience"),
    # Technical projects
    (r"(technical\s+)?projects?(\s+&\s+.+)?", "technical_projects"),
    (r"side\s+projects?", "technical_projects"),
    (r"personal\s+projects?", "technical_projects"),
    # Education
    (r"education(\s+&\s+.+)?", "education"),
    (r"academic\s+background", "education"),
    (r"qualifications?", "education"),
    # Skills
    (r"(technical\s+)?skills?(\s+&\s+.+)?", "skills"),
    (r"competenc(e|ies)", "skills"),
    (r"technologies(\s+&\s+.+)?", "skills"),
    (r"tools(\s+&\s+.+)?", "skills"),
    # Summaries
    (r"(professional\s+)?summary", "summary"),
    (r"(personal\s+)?(profile|statement|objective)", "summary"),
    (r"about\s+me", "summary"),
    # Cover letter
    (r"cover\s+letter", "cover_letter"),
    (r"application\s+letter", "cover_letter"),
]

_compiled: list[tuple[re.Pattern[str], SectionType]] = [
    (re.compile(pat, re.IGNORECASE), stype)
    for pat, stype in _HEADING_PATTERNS
]


@dataclass
class RawSection:
    heading: str
    raw_text: str
    detected_type: SectionType
    confidence: float                # 0.0–1.0
    warnings: list[str] = field(default_factory=list)


def _classify_heading(text: str) -> tuple[SectionType, float]:
    """Return (detected_type, confidence) for a heading string."""
    stripped = text.strip().lower()
    for pattern, stype in _compiled:
        if pattern.fullmatch(stripped):
            return stype, 1.0
        if pattern.search(stripped):
            return stype, 0.75
    return "other", 0.3


def _is_heading_style(style_name: str) -> bool:
    return bool(re.match(r"heading\s+\d", style_name, re.IGNORECASE))


def _is_bold_uppercase(para) -> bool:
    """Return True if the paragraph looks like a manual heading (bold and/or ALL CAPS)."""
    text = para.text.strip()
    if not text or len(text) > 80:
        return False
    is_upper = text == text.upper() and any(c.isalpha() for c in text)
    is_bold = bool(para.runs and para.runs[0].bold)
    return is_upper or is_bold


# -----------------------  DOCX extraction  -----------------------

def extract_docx_sections(path: Path) -> list[RawSection]:
    """Parse a .docx file and return a list of RawSection objects.

    Section breaks are detected from:
    1. Paragraph styles named ``Heading N``
    2. Paragraphs whose first run is bold or whose text is ALL CAPS <= 80 chars
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX parsing") from None

    doc = Document(str(path))

    sections: list[RawSection] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def _flush() -> None:
        if current_heading is None:
            return
        raw = "\n".join(current_lines).strip()
        stype, conf = _classify_heading(current_heading)
        sections.append(
            RawSection(
                heading=current_heading,
                raw_text=raw,
                detected_type=stype,
                confidence=conf,
            )
        )

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if _is_heading_style(para.style.name) or _is_bold_uppercase(para):
            _flush()
            current_heading = text
            current_lines = []
        else:
            if current_heading is None:
                current_heading = "Profile"
            current_lines.append(text)

    _flush()

    # If nothing was detected, return the whole document as one "other" section
    if not sections:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        sections.append(
            RawSection(
                heading="Full Document",
                raw_text=full_text,
                detected_type="other",
                confidence=0.2,
                warnings=["No headings detected; whole document returned as one section"],
            )
        )

    return sections


# -----------------------  PDF extraction  -----------------------

_MIN_HEADING_FONT_SIZE = 12.0  # points — larger than body text


def _median_font_size(words: list[dict]) -> float:
    sizes = sorted(w.get("size", 10) for w in words if w.get("size"))
    if not sizes:
        return 10.0
    return sizes[len(sizes) // 2]


def extract_pdf_sections(path: Path) -> list[RawSection]:
    """Parse a .pdf file using pdfplumber and return RawSection objects.

    Headings are detected by font-size: lines whose characters are
    significantly larger than the body text font OR in ALL CAPS are
    treated as section headers.
    """
    try:
        import pdfplumber  # type: ignore[import-untyped]
    except ImportError:
        raise RuntimeError(
            "pdfplumber is required for PDF parsing. "
            "Run: uv add pdfplumber"
        ) from None

    all_lines: list[dict] = []  # {"text": str, "size": float, "body_size": float}

    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            words = page.extract_words(extra_attrs=["size"])
            if not words:
                continue
            body_size = _median_font_size(words)

            # Group words into lines by approximate y-position (bucketed to 2px)
            lines_by_y: dict[int, list[dict]] = {}
            for w in words:
                y_bucket = round(w.get("top", 0) / 2) * 2
                lines_by_y.setdefault(y_bucket, []).append(w)

            for y_bucket in sorted(lines_by_y):
                line_words = sorted(lines_by_y[y_bucket], key=lambda w: w.get("x0", 0))
                line_text = " ".join(w["text"] for w in line_words).strip()
                avg_size = sum(w.get("size", body_size) for w in line_words) / len(line_words)
                all_lines.append({"text": line_text, "size": avg_size, "body_size": body_size})

    # Pass 2: split into sections at heading lines
    sections: list[RawSection] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    def _flush() -> None:
        if current_heading is None:
            return
        raw = "\n".join(current_lines).strip()
        stype, conf = _classify_heading(current_heading)
        sections.append(
            RawSection(
                heading=current_heading,
                raw_text=raw,
                detected_type=stype,
                confidence=conf,
            )
        )

    for entry in all_lines:
        text = entry["text"]
        avg_size = entry["size"]
        body_size = entry["body_size"]

        is_larger = avg_size >= max(body_size * 1.1, _MIN_HEADING_FONT_SIZE)
        is_caps_short = (
            text == text.upper()
            and any(c.isalpha() for c in text)
            and len(text) <= 60
        )

        if is_larger or is_caps_short:
            _flush()
            current_heading = text
            current_lines = []
        else:
            if current_heading is None:
                current_heading = "Profile"
            current_lines.append(text)

    _flush()

    if not sections:
        full_text = "\n".join(e["text"] for e in all_lines)
        sections.append(
            RawSection(
                heading="Full Document",
                raw_text=full_text,
                detected_type="other",
                confidence=0.2,
                warnings=["No headings detected; whole document returned as one section"],
            )
        )

    return sections


# -----------------------  unified entry point  -----------------------

def extract_sections(path: Path) -> list[RawSection]:
    """Auto-dispatch to DOCX or PDF extractor based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_sections(path)
    elif suffix == ".pdf":
        return extract_pdf_sections(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix!r}. Use .docx or .pdf")


def sections_to_json(sections: list[RawSection]) -> list[dict]:
    """Serialise RawSection list to JSON-safe dicts for the API response."""
    return [
        {
            "heading": s.heading,
            "raw_text": s.raw_text,
            "detected_type": s.detected_type,
            "confidence": s.confidence,
            "warnings": s.warnings,
        }
        for s in sections
    ]
